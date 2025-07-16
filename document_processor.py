import os
import io
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from docx import Document
from unstructured.partition.auto import partition
from langchain.docstore.document import Document as LangchainDocument
import re

def process_file(file_path: str, filename: str, metadata: Dict[str, Any]) -> List[LangchainDocument]:
    """
    Обрабатывает файл по пути, извлекает текст и разбивает на чанки.
    Поддерживает PDF, DOCX, TXT.
    """
    print(f"ℹ️ Начинаем обработку файла: {filename}")
    file_extension = os.path.splitext(filename)[1].lower()
    text = ""

    # Извлечение текста
    if file_extension == '.pdf':
        with open(file_path, "rb") as file:
            reader = PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
    elif file_extension == '.docx':
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + " "
    elif file_extension == '.txt':
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()
    else:
        # Использование unstructured для других форматов
        elements = partition(file_path)
        text = "\n\n".join([str(el) for el in elements])

    if not text:
        print("❌ Не удалось извлечь текст из документа.")
        return []

    # Простая стратегия разбиения на чанки
    chunks = re.split(r'(?<=[.!?])\s+', text)
    processed_chunks = []

    current_chunk = ""
    for chunk in chunks:
        if len(current_chunk) < 500:
            current_chunk += " " + chunk
        else:
            processed_chunks.append(current_chunk.strip())
            current_chunk = chunk
    if current_chunk:
        processed_chunks.append(current_chunk.strip())

    # Преобразование в формат LangChain Document
    documents = [
        LangchainDocument(page_content=chunk, metadata=metadata)
        for chunk in processed_chunks
    ]

    print(f"✅ Файл обработан. Извлечено {len(documents)} чанков.")
    return documents

def process_uploaded_file(file_stream: io.BytesIO, filename: str, metadata: Dict[str, Any]) -> List[LangchainDocument]:
    """
    Обрабатывает файл из потока, извлекает текст и разбивает на чанки.
    Используется для файлов, загруженных через API.
    """
    print(f"ℹ️ Начинаем обработку загруженного файла: {filename}")

    temp_path = f"./temp_{filename}"
    with open(temp_path, "wb") as temp_file:
        temp_file.write(file_stream.getvalue())

    try:
        documents = process_file(temp_path, filename, metadata)
    finally:
        os.remove(temp_path)

    return documents
